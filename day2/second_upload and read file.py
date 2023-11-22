import streamlit as st
import pandas as pd
from io import StringIO
import docx
import olefile
import zlib
import struct

st.subheader("csv, xlsx파일")
uploaded_file = st.file_uploader("csv/xlsx 파일을 업로드해주세요. ", type = ['csv', 'xlsx'])
if uploaded_file is not None:
    if uploaded_file.name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
    elif uploaded_file.name.endswith("xlsx"):
        df = pd.read_excel(uploaded_file)
    else:
        st.error("지원하지 않는 파일입니다. ")
    st.write(df.head(5))

st.subheader("DOCX 파일")
uploaded_file = st.file_uploader("word 파일을 업로드해주세요. ", type = 'docx')
if uploaded_file is not None:
    doc = docx.Document(uploaded_file)
    st.write("### 본문 텍스트")
    for par in doc.paragraphs:
        st.write(par.text)

    st.write("### 표")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    st.write(par.text)
            
st.subheader("한글 파일_1")
uploaded_file = st.file_uploader("한글 파일을 업로드해주세요.(1장짜리) ", type = 'hwp')
if uploaded_file is not None:
    hwp = olefile.OleFileIO(uploaded_file)
    encoded_text = hwp.openstream('PrvText').read()
    decoded_text = encoded_text.decode('UTF-16')
    st.write(decoded_text)



# 한글 파일 읽어오기 함수 정의
def get_hwp_text(filename):
    f = olefile.OleFileIO(filename)
    dirs = f.listdir()

    # HWP 파일 검증
    if ["FileHeader"] not in dirs or \
       ["\x05HwpSummaryInformation"] not in dirs:
        raise Exception("Not Valid HWP.")

    # 문서 포맷 압축 여부 확인
    header = f.openstream("FileHeader")
    header_data = header.read()
    is_compressed = (header_data[36] & 1) == 1

    # Body Sections 불러오기
    nums = []
    for d in dirs:
        if d[0] == "BodyText":
            nums.append(int(d[1][len("Section"):]))
    sections = ["BodyText/Section"+str(x) for x in sorted(nums)]

    # 전체 text 추출
    text = ""
    for section in sections:
        bodytext = f.openstream(section)
        data = bodytext.read()
        if is_compressed:
            unpacked_data = zlib.decompress(data, -15)
        else:
            unpacked_data = data
    
        # 각 Section 내 text 추출    
        section_text = ""
        i = 0
        size = len(unpacked_data)
        while i < size:
            header = struct.unpack_from("<I", unpacked_data, i)[0]
            rec_type = header & 0x3ff
            rec_len = (header >> 20) & 0xfff

            if rec_type in [67]:
                rec_data = unpacked_data[i+4:i+4+rec_len]
                section_text += rec_data.decode('utf-16')
                section_text += "\n"

            i += 4 + rec_len

        text += section_text
        text += "\n"

    return text


st.subheader("한글 파일_2")
uploaded_file = st.file_uploader("한글 파일을 업로드해주세요.(2페이지 이상) ", type = 'hwp')
if uploaded_file is not None:
    hwp_text = get_hwp_text(uploaded_file)
    st.write(hwp_text)